"""テンプレート管理モジュール（修正版）"""
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class TemplateManager:
    def __init__(self, templates):
        self.templates = templates
    
    def apply_template(self, doc, template_name, **kwargs):
        """テンプレートを文書に適用"""
        if template_name not in self.templates:
            raise ValueError(f"テンプレート '{template_name}' が見つかりません")
        
        template = self.templates[template_name]
        
        # セクションを取得
        if not doc.sections:
            return doc
        
        section = doc.sections[0]
        
        # ヘッダー設定
        if template.get('header'):
            try:
                header = section.header
                if not header.paragraphs:
                    header_para = header.add_paragraph()
                else:
                    header_para = header.paragraphs[0]
                
                header_para.text = template['header'].format(**kwargs)
                header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                if header_para.runs:
                    for run in header_para.runs:
                        run.font.size = Pt(10)
            except Exception as e:
                print(f"ヘッダー設定エラー: {e}")
        
        # フッター設定
        if template.get('footer'):
            try:
                footer = section.footer
                if not footer.paragraphs:
                    footer_para = footer.add_paragraph()
                else:
                    footer_para = footer.paragraphs[0]
                
                footer_para.text = template['footer'].format(**kwargs)
                footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                if footer_para.runs:
                    for run in footer_para.runs:
                        run.font.size = Pt(10)
            except Exception as e:
                print(f"フッター設定エラー: {e}")
        
        return doc
    
    def list_templates(self):
        """利用可能なテンプレート一覧"""
        return {name: tmpl['name'] for name, tmpl in self.templates.items()}
