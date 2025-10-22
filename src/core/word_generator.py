"""Word文書生成モジュール（サイズ調整版）"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from typing import List, Dict, Any
import re

class WordGenerator:
    def __init__(self, style_config: dict, math_converter):
        self.config = style_config
        self.math_converter = math_converter
        self.problem_counter = 0
    
    def create_document(self, problems: List[Dict[str, Any]]) -> Document:
        """複数の問題からWord文書を生成"""
        doc = Document()
        self._apply_global_style(doc)
        
        for i, problem in enumerate(problems):
            self.problem_counter = i + 1
            self._add_problem(doc, problem)
            
            # 大問間の間隔（最後の問題以外）
            if i < len(problems) - 1:
                doc.add_paragraph()
        
        return doc
    
    def _apply_global_style(self, doc: Document):
        """文書全体のスタイル設定"""
        sections = doc.sections
        for section in sections:
            # 余白設定
            margin = Cm(self.config['page_margin'])
            section.top_margin = margin
            section.bottom_margin = margin
            section.left_margin = margin
            section.right_margin = margin
        
        # スタイル定義
        styles = doc.styles
        
        # 見出しスタイル
        if 'Heading 1' in styles:
            style = styles['Heading 1']
            style.font.name = self.config['title_font']
            style.font.size = Pt(self.config['title_size'])
            style.font.bold = self.config['title_bold']
    
    def _add_problem(self, doc: Document, problem: Dict[str, Any]):
        """問題を文書に追加"""
        # タイトル
        if problem['title']:
            title = doc.add_heading(problem['title'], level=1)
            self._apply_title_style(title)
        else:
            title = doc.add_heading(f'大問{self.problem_counter}', level=1)
            self._apply_title_style(title)
        
        # 問題文（テキストと数式が混在）
        current_para = None
        
        for element in problem['text']:
            if element['type'] == 'text':
                # テキスト要素
                if current_para is None:
                    current_para = doc.add_paragraph()
                    self._apply_body_style(current_para)
                
                # テキストを追加
                run = current_para.add_run(element['content'])
                
            elif element['type'] == 'math':
                # インライン数式
                if current_para is None:
                    current_para = doc.add_paragraph()
                    self._apply_body_style(current_para)
                
                # 数式画像を追加（小さめに）
                try:
                    img_stream = self.math_converter.latex_to_image(element['content'])
                    run = current_para.add_run()
                    # インライン数式の高さを設定から取得
                    inline_height = self.config.get('inline_math_height', 14)
                    run.add_picture(img_stream, height=Pt(inline_height))
                except Exception as e:
                    # エラー時はテキストで表示
                    run = current_para.add_run(f"[{element['content']}]")
                    run.font.color.rgb = RGBColor(255, 0, 0)
        
        # 段落が作成されていれば、終了
        if current_para is not None:
            current_para = None
        
        # 独立した数式
        for equation in problem['equations']:
            self._add_display_math(doc, equation)
        
        # 選択肢
        if problem['choices']:
            self._add_choices(doc, problem['choices'])
    
    def _apply_title_style(self, paragraph):
        """タイトルスタイルを適用"""
        for run in paragraph.runs:
            run.font.name = self.config['title_font']
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.config['title_font'])
            run.font.size = Pt(self.config['title_size'])
            run.font.bold = self.config['title_bold']
    
    def _apply_body_style(self, paragraph):
        """本文スタイルを適用"""
        paragraph.paragraph_format.line_spacing = self.config['paragraph_spacing']
        
        for run in paragraph.runs:
            run.font.name = self.config['body_font']
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.config['body_font'])
            run.font.size = Pt(self.config['body_size'])
    
    def _add_display_math(self, doc: Document, latex_str: str):
        """独立した数式を追加（中央揃え）"""
        para = doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = para.add_run()
        
        try:
            # 数式を画像化
            img_stream = self.math_converter.latex_to_image(latex_str)
            # ディスプレイ数式の幅を設定から取得
            display_width = self.config.get('display_math_width', 2.5)
            run.add_picture(img_stream, width=Inches(display_width))
        except Exception as e:
            # エラー時
            run.text = f"[数式エラー: {latex_str[:30]}...]"
            run.font.color.rgb = RGBColor(255, 0, 0)
        
        para.paragraph_format.space_after = Pt(12)
    
    def _add_choices(self, doc: Document, choices: List[str]):
        """選択肢を追加"""
        for i, choice in enumerate(choices, 1):
            # 選択肢にもインライン数式が含まれる可能性がある
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(self.config['choice_indent'])
            
            # 番号を追加
            run = para.add_run(f'{i}. ')
            run.font.name = self.config['choice_font']
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.config['choice_font'])
            run.font.size = Pt(self.config['choice_size'])
            
            # 選択肢テキストを解析してインライン数式を処理
            self._add_text_with_inline_math(para, choice)
    
    def _add_text_with_inline_math(self, paragraph, text: str):
        """テキスト内のインライン数式を処理して段落に追加"""
        # インライン数式パターン: \(...\) または $...$
        pattern = r'\\\((.+?)\\\)|\$([^$]+)\$'
        
        last_end = 0
        for match in re.finditer(pattern, text):
            # 数式の前のテキスト
            if match.start() > last_end:
                plain_text = text[last_end:match.start()]
                run = paragraph.add_run(plain_text)
                run.font.name = self.config['choice_font']
                run._element.rPr.rFonts.set(qn('w:eastAsia'), self.config['choice_font'])
                run.font.size = Pt(self.config['choice_size'])
            
            # 数式部分
            latex_str = match.group(1) or match.group(2)
            if latex_str:
                try:
                    img_stream = self.math_converter.latex_to_image(latex_str.strip())
                    run = paragraph.add_run()
                    # 選択肢内のインライン数式も小さめに
                    inline_height = self.config.get('inline_math_height', 14)
                    run.add_picture(img_stream, height=Pt(inline_height))
                except Exception as e:
                    # エラー時はテキストで表示
                    run = paragraph.add_run(f"[{latex_str}]")
                    run.font.color.rgb = RGBColor(255, 0, 0)
            
            last_end = match.end()
        
        # 残りのテキスト
        if last_end < len(text):
            plain_text = text[last_end:]
            run = paragraph.add_run(plain_text)
            run.font.name = self.config['choice_font']
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.config['choice_font'])
            run.font.size = Pt(self.config['choice_size'])
